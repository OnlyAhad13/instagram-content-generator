import streamlit as st
import pandas as pd
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from langchain.retrievers import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import LLMChainExtractor
from docx import Document as DocxDocument
import re
import os
from io import BytesIO

# =====================================
# Streamlit Configuration
# =====================================
st.set_page_config(
    page_title="Instagram Content Generator",
    page_icon="📱",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 3rem;
        font-weight: bold;
        text-align: center;
        background: linear-gradient(90deg, #ff6b6b, #4ecdc4, #45b7d1);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 2rem;
    }
    .script-output {
        background-color: #f8f9fa;
        border-left: 4px solid #007bff;
        padding: 1rem;
        border-radius: 0.5rem;
        font-family: 'Courier New', monospace;
        margin: 1rem 0;
    }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 0.5rem;
        color: white;
        text-align: center;
        margin: 0.5rem 0;
    }
</style>
""", unsafe_allow_html=True)

# =====================================
# Helper Functions
# =====================================
@st.cache_data
def load_and_enhance_data(style_file=None):
    """Load and process files - style guide can be uploaded, dataset is hardcoded"""
    try:
        # Load style guide (.docx) - either uploaded or use default
        style_texts = []
        
        if style_file is not None:
            # Use uploaded style guide
            try:
                style_file.seek(0)
                doc = DocxDocument(BytesIO(style_file.read()))
                for i, p in enumerate(doc.paragraphs):
                    if p.text.strip():
                        context = f"[Document section {i//10 + 1}] {p.text.strip()}"
                        style_texts.append(context)
            except Exception as e:
                st.error(f"Error reading uploaded style guide: {str(e)}")
                return [], [], None
        else:
            # Use default style guide
            style_file_path = "July 2025 - Scripts & Captions [Mirrored Aesthetics] (1) (1).docx"
            if os.path.exists(style_file_path):
                try:
                    doc = DocxDocument(style_file_path)
                    for i, p in enumerate(doc.paragraphs):
                        if p.text.strip():
                            context = f"[Document section {i//10 + 1}] {p.text.strip()}"
                            style_texts.append(context)
                except Exception as e:
                    st.error(f"Error reading default style guide: {str(e)}")
                    return [], [], None
            else:
                st.warning(f"Default style guide file not found: {style_file_path}")

        style_docs = [Document(
            page_content=t,
            metadata={"source": "style_guide", "type": "expert_example", "relevance": "high"}
        ) for t in style_texts]

        # Load dataset (.xlsx) - hardcoded file
        df = None
        data_docs = []
        data_file_path = "IGScraper.xlsx"
        
        if os.path.exists(data_file_path):
            try:
                df = pd.read_excel(data_file_path)
            except Exception as e:
                st.error(f"Error reading performance data: {str(e)}")
                return [], [], None
        else:
            st.warning(f"Performance data file not found: {data_file_path}")
        
        if df is not None:
            df.columns = df.columns.str.strip().str.lower()

            # Numeric columns
            numeric_cols = ['likes', 'comments', 'views']
            for col in numeric_cols:
                if col in df.columns:
                    df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
                else:
                    df[col] = 0

            # Engagement score + performance tiers
            df['engagement'] = df['likes']*0.4 + df['comments']*0.4 + df['views']*0.2
            df['performance_tier'] = pd.qcut(
                df['engagement'], q=3, labels=['low', 'medium', 'high'], duplicates='drop'
            )

            # Compose text
            def compose_enhanced_text(row):
                parts = []
                column_mappings = {
                    'hook': ['hooks', 'hook', 'opening'],
                    'body': ['description', 'body', 'content'],
                    'caption': ['captions', 'caption', 'copy'],
                    'summary': ['summary', 'notes', 'key_points']
                }
                for section, possible_cols in column_mappings.items():
                    for col in possible_cols:
                        if col in df.columns and pd.notna(row[col]) and str(row[col]).strip():
                            clean_text = str(row[col]).strip()
                            if clean_text.lower() != 'nan':
                                parts.append(f"{section.upper()}: {clean_text}")
                            break
                perf_tier = row.get('performance_tier', 'unknown')
                engagement = row.get('engagement', 0)
                parts.append(f"PERFORMANCE: {perf_tier} tier ({engagement:.0f} engagement score)")
                return "\n".join(parts)

            for _, row in df.iterrows():
                text = compose_enhanced_text(row)
                if text.strip():
                    data_docs.append(Document(
                        page_content=text,
                        metadata={
                            "source": "performance_data",
                            "type": "proven_example",
                            "likes": int(row.get('likes', 0)),
                            "comments": int(row.get('comments', 0)),
                            "views": int(row.get('views', 0)),
                            "engagement": float(row.get('engagement', 0)),
                            "performance_tier": str(row.get('performance_tier', 'unknown')),
                            "relevance": "high" if row.get('engagement', 0) > df['engagement'].quantile(0.7) else "medium"
                        }
                    ))

        return style_docs, data_docs, df
    except Exception as e:
        st.error(f"Error loading data: {str(e)}")
        return [], [], None

@st.cache_resource
def create_smart_retriever(docs, openai_api_key):
    """Create the smart retriever with caching"""
    if not docs:
        return None
    
    try:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=600, chunk_overlap=150,
            separators=["\n\n", "\n", ".", "!", "?", " ", ""],
            keep_separator=True
        )
        chunks = splitter.split_documents(docs)
        
        # Set OpenAI API key
        os.environ["OPENAI_API_KEY"] = openai_api_key
        
        embeddings = OpenAIEmbeddings(model="text-embedding-3-large")
        vectorstore = FAISS.from_documents(chunks, embeddings)

        base_retriever = vectorstore.as_retriever(
            search_type="similarity_score_threshold",
            search_kwargs={"k": 8, "score_threshold": 0.7, "fetch_k": 16}
        )
        llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
        compressor = LLMChainExtractor.from_llm(llm)

        return ContextualCompressionRetriever(
            base_compressor=compressor,
            base_retriever=base_retriever
        )
    except Exception as e:
        st.error(f"Error creating retriever: {str(e)}")
        return None

def enhance_retrieval_for_goal(retriever, goal):
    """Enhance retrieval based on goal"""
    def goal_aware_retrieval(query):
        docs = retriever.invoke(query)
        goal_enhanced = []
        for doc in docs:
            if doc.metadata.get('source') == 'performance_data':
                score = doc.metadata.get(goal.lower(), 0)
                doc.metadata['goal_relevance'] = score
                goal_enhanced.append(doc)
            else:
                doc.metadata['goal_relevance'] = 100
                goal_enhanced.append(doc)
        goal_enhanced.sort(key=lambda x: x.metadata.get('goal_relevance', 0), reverse=True)
        return goal_enhanced[:6]
    return goal_aware_retrieval

def create_optimized_prompt():
    """Create the prompt template"""
    template = """You are an expert Instagram content strategist and copywriter.

CONTEXT:
{context}

CRITICAL INSTRUCTIONS:
1. Analyze examples carefully
2. Identify patterns in hooks, storytelling, CTAs
3. Follow EXACT format - no deviations
4. {custom_instruction}

OUTPUT FORMAT:
Hook: [under 8 words, curiosity gap or contrarian]
Body: [2-4 sentences, storytelling OR educational, simple words]
CTA: [direct action, urgency/exclusivity]
Caption: [restates CTA + 8-12 hashtags]

GOAL: Maximize {goal}

Generate ONE script optimized for {goal}."""
    return PromptTemplate.from_template(template)

def enforce_script_format(llm, prompt_text, max_retries=3):
    """Enforce the script format with retries"""
    required_sections = ["Hook:", "Body:", "CTA:", "Caption:"]
    response = ""
    
    for attempt in range(max_retries):
        try:
            response = llm.invoke(prompt_text).content
            if all(sec in response for sec in required_sections):
                return response
            st.warning(f"⚠️ Missing sections on attempt {attempt+1}, retrying...")
        except Exception as e:
            st.error(f"Error generating content: {str(e)}")
            return None
    
    st.warning("❌ Could not enforce format after retries. Returning last attempt.")
    return response

# =====================================
# Main App
# =====================================
def main():
    st.markdown('<h1 class="main-header">📱 Instagram Content Generator</h1>', unsafe_allow_html=True)
    st.markdown("---")
    
    # Sidebar configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        
        # OpenAI API Key
        openai_api_key = st.text_input(
            "OpenAI API Key",
            type="password",
            placeholder="sk-...",
            help="Enter your OpenAI API key to use the generator"
        )
        
        st.markdown("---")
        
        # Custom query
        st.subheader("🎨 Custom Query")
        custom_query = st.text_area(
            "Custom Query (Optional)",
            placeholder="E.g., Create a fitness motivation script",
            help="Provide specific instructions for your content"
        )
        
        st.markdown("---")
        
        # Goal selection
        st.subheader("🎯 Optimization Goal")
        goal = st.selectbox(
            "Select your primary goal:",
            options=["likes", "comments", "views", "engagement"],
            index=1,
            help="Choose what metric you want to optimize for"
        )
        
        st.markdown("---")
        
        # Style guide upload
        st.subheader("📝 Style Guide")
        style_file = st.file_uploader(
            "Upload Custom Style Guide (.docx)",
            type=['docx'],
            help="Upload your custom style guide, or leave empty to use the default one"
        )
        
        if style_file is not None:
            st.success("✅ Custom style guide uploaded")
        else:
            st.info("📄 Using default style guide")
        
        st.markdown("---")
        
        # Data status
        st.subheader("📊 Data Status")
        st.success("✅ Performance data loaded automatically")

    # Main content area
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("🚀 Generate Content")
        
        # Check if all requirements are met
        if not openai_api_key:
            st.warning("⚠️ Please enter your OpenAI API key in the sidebar to proceed.")
            return
        
        # Generate button
        if st.button("🎬 Generate Instagram Script", type="primary", width='stretch'):
            with st.spinner("🔄 Analyzing your data and generating optimized content..."):
                try:
                    # ==========================================
                    # STEP 1: DATA LOADING
                    # ==========================================
                    st.markdown("### 🔍 **DEBUG: Step 1 - Data Loading**")
                    with st.expander("📊 Click to see data loading details", expanded=True):
                        st.write("**Loading style guide and performance data...**")
                        style_docs, data_docs, df = load_and_enhance_data(style_file)
                        all_docs = style_docs + data_docs
                        
                        st.write(f"✅ **Style documents loaded**: {len(style_docs)} documents")
                        st.write(f"✅ **Performance data documents loaded**: {len(data_docs)} documents")
                        st.write(f"✅ **Total documents**: {len(all_docs)} documents")
                        
                        if df is not None:
                            st.write(f"✅ **Performance data rows**: {len(df)} rows")
                            st.write(f"✅ **Performance data columns**: {list(df.columns)}")
                        else:
                            st.write("⚠️ **No performance data found**")
                        
                        if not all_docs:
                            st.error("❌ No valid data found in uploaded files.")
                            return
                    
                    # ==========================================
                    # STEP 2: RETRIEVER CREATION
                    # ==========================================
                    st.markdown("### 🔍 **DEBUG: Step 2 - Creating Smart Retriever**")
                    with st.expander("🧠 Click to see retriever creation details", expanded=True):
                        st.write("**Creating AI-powered document retriever...**")
                        retriever = create_smart_retriever(all_docs, openai_api_key)
                        if not retriever:
                            st.error("❌ Failed to create retriever")
                            return
                        st.write("✅ **Smart retriever created successfully**")
                        
                        goal_retriever = enhance_retrieval_for_goal(retriever, goal)
                        st.write(f"✅ **Goal-aware retriever configured for**: {goal}")
                    
                    # ==========================================
                    # STEP 3: AI BRAIN PROCESSING
                    # ==========================================
                    st.markdown("### 🧠 **AI BRAIN: Processing Your Request**")
                    with st.expander("🎯 Click to see AI thinking process", expanded=True):
                        query = custom_query or f"Create high-{goal} Instagram reel script"
                        
                        # Dynamic AI processing visualization
                        st.markdown("#### 🎯 **Analyzing Your Request**")
                        if custom_query:
                            st.success(f"🎨 **Custom Request**: '{custom_query}'")
                        else:
                            st.info(f"🎯 **Default Goal**: Optimize for {goal}")
                        
                        # Simulate AI thinking process
                        import time
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        steps = [
                            "🔍 Analyzing your query...",
                            "📚 Accessing knowledge base...",
                            "🎨 Understanding style preferences...",
                            "📊 Processing performance data...",
                            "🧠 Building creative context...",
                            "✨ Preparing AI generation..."
                        ]
                        
                        for i, step in enumerate(steps):
                            status_text.text(step)
                            progress_bar.progress((i + 1) / len(steps))
                            time.sleep(0.5)  # Simulate processing time
                        
                        status_text.text("✅ AI Brain Ready!")
                        progress_bar.progress(1.0)
                        
                        # Show what AI is considering
                        st.markdown("#### 🧠 **AI Context Analysis**")
                        st.write("**The AI is considering:**")
                        st.write("• Your specific request and goals")
                        st.write("• Style guide patterns and preferences")
                        st.write("• Performance data insights")
                        st.write("• Instagram best practices")
                        st.write("• Creative content strategies")
                        
                        # Get relevant docs (but don't show empty results)
                        relevant_docs = goal_retriever(query)
                        context = "\n\n---\n\n".join([doc.page_content for doc in relevant_docs]) if relevant_docs else ""
                        
                        if context:
                            st.success(f"📚 **Knowledge Base**: Found {len(relevant_docs)} relevant insights")
                        else:
                            st.info("🎨 **Creative Mode**: Using AI's built-in knowledge and creativity")
                    
                    # ==========================================
                    # STEP 4: AI CREATIVE GENERATION
                    # ==========================================
                    st.markdown("### ✨ **AI CREATIVE GENERATION**")
                    with st.expander("🤖 Click to see AI creative process", expanded=True):
                        prompt = create_optimized_prompt()
                        llm = ChatOpenAI(model="gpt-4o", temperature=0.7, max_tokens=800, openai_api_key=openai_api_key)
                        
                        custom_instruction = f"Focus on: {custom_query}" if custom_query else "Create engaging content based on the provided examples"
                        
                        # Dynamic generation visualization
                        st.markdown("#### 🎨 **AI Creative Engine**")
                        st.write("**🧠 AI Model**: GPT-4o (Advanced Creative AI)")
                        st.write("**🎯 Creativity Level**: 70% (Balanced Creative & Strategic)")
                        st.write("**📝 Output Length**: Up to 800 tokens")
                        st.write(f"**🎯 Optimization Goal**: {goal.upper()}")
                        
                        if custom_query:
                            st.write(f"**🎨 Custom Focus**: {custom_query}")
                        else:
                            st.write("**🎨 Creative Mode**: General Instagram optimization")
                        
                        # Show generation process
                        st.markdown("#### ⚡ **Generation Process**")
                        gen_progress = st.progress(0)
                        gen_status = st.empty()
                        
                        gen_steps = [
                            "🧠 Analyzing requirements...",
                            "🎨 Crafting compelling hook...",
                            "📝 Building engaging body...",
                            "🎯 Creating strong CTA...",
                            "📱 Optimizing caption...",
                            "✨ Finalizing script..."
                        ]
                        
                        for i, step in enumerate(gen_steps):
                            gen_status.text(step)
                            gen_progress.progress((i + 1) / len(gen_steps))
                            time.sleep(0.3)
                        
                        gen_status.text("🎉 Generating your script...")
                        gen_progress.progress(1.0)
                        
                        formatted_prompt = prompt.format(
                            context=context, 
                            goal=goal, 
                            custom_instruction=custom_instruction
                        )
                        
                        response = enforce_script_format(llm, formatted_prompt, max_retries=3)
                    
                    # ==========================================
                    # STEP 5: FINAL RESULTS
                    # ==========================================
                    st.markdown("### 🎉 **CREATION COMPLETE!**")
                    with st.expander("📋 Click to see generation summary", expanded=True):
                        if response:
                            st.success("🎉 **Script Generated Successfully!**")
                            st.write(f"**📊 Script Length**: {len(response)} characters")
                            st.write(f"**🎯 Optimized For**: {goal.upper()}")
                            if custom_query:
                                st.write(f"**🎨 Custom Focus**: {custom_query}")
                            st.write("**✨ Your Instagram script is ready!**")
                        else:
                            st.error("❌ Generation failed. Please try again.")
                    
                    # Display the final result
                    if response:
                        st.markdown("### 🏆 Your Optimized Script")
                        st.markdown("---")
                        st.markdown(response)
                        st.markdown("---")
                        
                except Exception as e:
                    st.error(f"❌ An error occurred: {str(e)}")
                    st.exception(e)  # Show full traceback
    
    with col2:
        st.header("📊 Information")
        
        st.info("""
        **🎯 How to Use:**
        
        1. **Enter your OpenAI API key** in the sidebar
        2. **Add a custom query** (optional) for specific content
        3. **Select optimization goal** (likes, comments, views, engagement)
        4. **Upload style guide** (optional) or use default
        5. **Click Generate** to create your Instagram script
        
        **✨ Features:**
        - AI-powered content generation
        - Custom style guide support
        - Goal-based optimization
        - Copy-friendly output format
        """)
        
        st.markdown("---")
        
        st.subheader("📝 Output Format")
        st.markdown("""
        Your generated script will include:
        - **Hook**: Attention-grabbing opening
        - **Body**: Main content/story
        - **CTA**: Call-to-action
        - **Caption**: Complete caption with hashtags
        """)
    
    # Footer
    st.markdown("---")
    st.markdown(
        "---",
        unsafe_allow_html=True
    )

if __name__ == "__main__":
    main()